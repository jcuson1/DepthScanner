from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUTPUT_NAME = "Инструкция_по_логам_HelloAR.docx"


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_subtitle(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(11)


def add_heading(document: Document, text: str) -> None:
    document.add_heading(text, level=1)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    output_path = root / OUTPUT_NAME

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    add_title(document, "Инструкция по логам HelloAR")
    add_subtitle(
        document,
        "Записка для команды. Что лежит в логах, как это читать и что именно уходит в share.",
    )
    add_paragraph(
        document,
        f"Актуально на {datetime.now().strftime('%d.%m.%Y %H:%M')}.",
    )

    add_heading(document, "Что это за файл")
    add_paragraph(
        document,
        "Ниже короткая рабочая инструкция по логам, которые пишет приложение HelloAR. "
        "Это не описание реализации, а памятка по содержимому: какие данные попадают в лог, "
        "какие секции за что отвечают, что именно формируется на старте, что дописывается по ходу "
        "работы и что отправляется через кнопку Share logs."
    )

    add_heading(document, "Где лежат логи")
    add_bullets(
        document,
        [
            "На каждый запуск создаётся отдельный текстовый файл.",
            "Имя файла: helloar-log-YYYYMMDD-HHMMSS.txt.",
            "Файл создаётся в каталоге logs внутри app external files, если он доступен.",
            "Если external files недоступен, используется внутренний каталог files/logs.",
            "Кнопка Share logs отдаёт наружу текущий активный лог-файл через FileProvider."
        ],
    )

    add_heading(document, "Как устроена строка лога")
    add_paragraph(
        document,
        "Базовый формат строки такой:"
    )
    add_paragraph(
        document,
        "YYYY-MM-DD HH:MM:SS.mmm | LEVEL | SECTION | MESSAGE"
    )
    add_paragraph(
        document,
        "Если к записи прикладывается исключение, сразу после основной строки идёт stack trace."
    )
    add_bullets(
        document,
        [
            "Время — локальная дата и время с миллисекундами.",
            "LEVEL — INFO, WARN, ERROR или LOGCAT.",
            "SECTION — блок, к которому относится сообщение.",
            "MESSAGE — полезная нагрузка: состояние, параметр, событие, причина ошибки или текст предупреждения."
        ],
    )

    add_heading(document, "Что лежит в шапке файла")
    add_paragraph(
        document,
        "В начале каждого файла приложение пишет служебный блок с разделителями "
        "\"================================================================\"."
    )
    add_bullets(
        document,
        [
            "название лога;",
            "время старта сессии;",
            "справочник по уровням INFO / WARN / ERROR / LOGCAT;",
            "список секций, которые дальше встречаются по файлу."
        ],
    )
    add_paragraph(
        document,
        "Этот блок нужен только для чтения файла человеком. В нём нет runtime-состояний, это именно "
        "служебная шапка, чтобы было понятно, как дальше ориентироваться по содержимому."
    )

    add_heading(document, "Какие секции есть в логах")
    add_bullets(
        document,
        [
            "SESSION — общая информация по текущей сессии приложения.",
            "DEVICE — железо и системные параметры телефона.",
            "GPU — графический стек после поднятия GL-контекста.",
            "ARCORE — доступность ARCore, создание session, configure, resume, pause, close.",
            "LIFECYCLE — события activity: onCreate, onResume, onPause, onDestroy.",
            "PERMISSION — состояние camera permission.",
            "RENDER — инициализация шейдеров, surface size, light estimate state.",
            "DEPTH — режим визуализации глубины и готовность depth-потоков.",
            "TRACKING — tracking state, failure reason, наличие отслеживаемых плоскостей.",
            "INPUT — меню, тапы, создание anchor, выбор настроек.",
            "SHARE — отправка лог-файла наружу.",
            "LOGCAT — строки logcat текущего процесса, если они доступны.",
            "CRASH — непойманное падение процесса."
        ],
    )

    add_heading(document, "Что приложение пишет на старте")
    add_paragraph(
        document,
        "Сразу после создания DiagnosticsLogger пишется стартовый блок SESSION и DEVICE."
    )
    add_bullets(
        document,
        [
            "Application started — факт старта сессии.",
            "Package — package name приложения.",
            "Version — versionName из package info.",
            "Manufacturer, Brand, Model — базовая марка устройства.",
            "Device, Product, Board, Hardware — низкоуровневые идентификаторы платформы.",
            "Fingerprint — отпечаток прошивки.",
            "Android — версия Android и SDK level.",
            "Supported ABIs — поддерживаемые архитектуры.",
            "SoC manufacturer / SoC model — если Android отдаёт эти поля.",
            "Memory class / Large memory class — лимиты памяти для приложения.",
            "Total memory / Available memory — объём памяти устройства на момент старта.",
            "OpenGL ES version — версия GL ES из DeviceConfigurationInfo.",
            "CPU info — первые строки /proc/cpuinfo."
        ],
    )

    add_heading(document, "Что приложение пишет по ARCore до создания session")
    add_paragraph(
        document,
        "До создания Session приложение вызывает checkAvailability и пишет результат в секцию ARCORE."
    )
    add_bullets(
        document,
        [
            "Availability: <status> — сырое значение из ArCoreApk.Availability.",
            "Support status: ARCore is installed and available — если статус SUPPORTED_INSTALLED.",
            "Support status: ARCore is not ready yet or unsupported — если статус неготовый или неподдерживаемый.",
            "ARCore installation requested by Google Play Services for AR — если requestInstall вернул INSTALL_REQUESTED.",
            "ARCore installation flow finished successfully — если установка завершилась и можно продолжать."
        ],
    )
    add_paragraph(
        document,
        "Если на этапе создания session ловится исключение, в лог пишется текст проблемы и полный stack trace. "
        "Там же остаётся человекочитаемый смысл ошибки: не установлен ARCore, пользователь отказался, "
        "слишком старая версия APK, старая версия SDK, устройство не поддерживается или произошёл общий сбой."
    )

    add_heading(document, "Что пишется по permission")
    add_bullets(
        document,
        [
            "Camera permission missing, requesting permission dialog — права ещё нет, система открывает запрос.",
            "Camera permission granted — пользователь выдал доступ.",
            "Camera permission denied — пользователь отказал.",
            "Camera permission denied with 'Do not ask again' — пользователь отказал и отключил повторный запрос."
        ],
    )

    add_heading(document, "Что пишется после поднятия OpenGL")
    add_paragraph(
        document,
        "После создания GL-ресурсов в секцию GPU пишутся три ключевых поля:"
    )
    add_bullets(
        document,
        [
            "Vendor — строка GL_VENDOR.",
            "Renderer — строка GL_RENDERER.",
            "Version — строка GL_VERSION."
        ],
    )
    add_paragraph(
        document,
        "В секцию RENDER при этом попадают сообщения о старте инициализации рендера, об успешной "
        "инициализации ресурсов и о каждом изменении размера surface."
    )

    add_heading(document, "Что пишется по session configure")
    add_paragraph(
        document,
        "После configureSession в ARCORE сохраняется уже готовая конфигурация session."
    )
    add_bullets(
        document,
        [
            "light=ENVIRONMENTAL_HDR;",
            "depth=<DepthMode>;",
            "instantPlacement=<InstantPlacementMode>;",
            "visualization=<DepthVisualizationMode>;",
            "occlusion=<true/false>."
        ],
    )
    add_paragraph(
        document,
        "Это текущая рабочая конфигурация на момент вызова session.configure(config)."
    )

    add_heading(document, "Что пишется во время работы кадра")
    add_paragraph(
        document,
        "В onDrawFrame логи не льются сплошным потоком на каждом кадре. Пишутся изменения состояния."
    )
    add_bullets(
        document,
        [
            "Visualization mode: <mode> — если переключили режим отображения глубины.",
            "Camera tracking state changed to <state> — если изменилось состояние tracking.",
            "Tracking failure reason: <reason> — если изменилась причина потери трекинга.",
            "Tracked plane available: <true/false> — если появился или исчез факт наличия отслеживаемой плоскости.",
            "Light estimate state: <state> — если изменилось состояние оценки освещения."
        ],
    )

    add_heading(document, "Что пишется по depth")
    add_paragraph(
        document,
        "В DEPTH идут две группы записей: выбранный режим и состояние доступности depth-потоков."
    )
    add_bullets(
        document,
        [
            "Visualization mode: CAMERA / FULL_DEPTH / RAW_DEPTH / CONFIDENCE / DEPTH_DOT_GRID / DISTANCE_PROBES.",
            "Full depth image is not available yet — full depth ещё не готов.",
            "Raw depth image is not available yet — raw depth ещё не готов.",
            "Raw depth confidence image is not available yet — confidence map ещё не готов.",
            "Showing first-run occlusion suggestion dialog — открыт диалог предложения включить depth occlusion.",
            "User enabled depth occlusion from suggestion dialog — пользователь включил occlusion.",
            "User declined depth occlusion from suggestion dialog — пользователь отказался."
        ],
    )

    add_heading(document, "Что пишется по tracking")
    add_paragraph(
        document,
        "В TRACKING лежат чистые служебные состояния ARCore, без пользовательских действий."
    )
    add_bullets(
        document,
        [
            "Tracking state камеры.",
            "Причина потери трекинга.",
            "Наличие отслеживаемой плоскости."
        ],
    )
    add_paragraph(
        document,
        "Если по логу видно, что tracking не дошёл до TRACKING или плоскости так и не появились, "
        "значит проблема не в тапе и не в anchor, а раньше — на уровне состояния самой AR-сцены."
    )

    add_heading(document, "Что пишется по пользовательским действиям")
    add_paragraph(
        document,
        "В INPUT попадают действия пользователя и всё, что идёт из UI."
    )
    add_bullets(
        document,
        [
            "Settings menu opened — открыто главное меню.",
            "Depth settings selected — пользователь зашёл в настройки Depth API.",
            "Depth visualization settings selected — пользователь зашёл в выбор режима визуализации.",
            "Instant placement settings selected — пользователь открыл настройки instant placement.",
            "Settings applied: ... — применён новый набор флагов и режимов.",
            "Tap detected at x=... y=... — координаты тапа.",
            "Anchor created from trackable type ... at pose ... — anchor создан, указан тип trackable и pose.",
            "Tap did not hit any ARCore trackable — тап не попал в пригодный объект.",
            "Anchor limit reached, removing oldest anchor — достигнут лимит anchor, старый anchor удалён."
        ],
    )

    add_heading(document, "Что пишется в LOGCAT")
    add_paragraph(
        document,
        "Если удалось поднять захват logcat по текущему pid, в файл добавляются строки вида:"
    )
    add_paragraph(
        document,
        "YYYY-MM-DD HH:MM:SS.mmm | LOGCAT | <исходная строка logcat>"
    )
    add_paragraph(
        document,
        "Это не замена основным структурированным логам, а дополнительный хвост процесса. "
        "Туда могут попасть сообщения приложения, Android runtime и части библиотек, если они "
        "видны в logcat текущего процесса."
    )

    add_heading(document, "Что пишется при падении")
    add_paragraph(
        document,
        "Если процесс падает через непойманное исключение, в секцию CRASH пишется:"
    )
    add_bullets(
        document,
        [
            "имя потока;",
            "текст исключения;",
            "полный stack trace."
        ],
    )

    add_heading(document, "Что уходит через Share logs")
    add_paragraph(
        document,
        "Кнопка Share logs берёт текущий активный log-файл и передаёт его во внешний share sheet."
    )
    add_bullets(
        document,
        [
            "Во внешний intent уходит сам файл через EXTRA_STREAM.",
            "В subject идёт строка Share diagnostics log.",
            "В text идёт строка Current log file is ready to send.",
            "Файл передаётся по URI через FileProvider с read permission."
        ],
    )
    add_paragraph(
        document,
        "То есть наружу уходит не выборка, не фрагмент и не последние N строк, а весь текущий файл "
        "целиком в том виде, в котором он на этот момент накопился."
    )

    add_heading(document, "Что делает Clear logs")
    add_bullets(
        document,
        [
            "останавливает текущий захват logcat;",
            "удаляет существующие лог-файлы из каталога logs;",
            "сбрасывает ссылку на текущий лог-файл;",
            "создаёт новый log-файл;",
            "снова пишет стартовую шапку;",
            "снова пишет SESSION и DEVICE;",
            "снова запускает logcat capture."
        ],
    )
    add_paragraph(
        document,
        "После очистки работа идёт уже в новом файле. Старое содержимое не переносится."
    )

    add_heading(document, "Быстрый способ читать лог")
    add_bullets(
        document,
        [
            "Сначала смотреть шапку и блок DEVICE — сразу видно, что за телефон и Android.",
            "Потом смотреть ARCORE — там видно availability, установку, создание session и configure.",
            "Дальше смотреть PERMISSION и LIFECYCLE — чтобы исключить банальные проблемы.",
            "Если проблема в глубине — идти в DEPTH.",
            "Если объект не ставится — смотреть TRACKING и INPUT рядом.",
            "Если был краш — сразу прыгать в CRASH и ERROR.",
            "Если своего сообщения не хватает — искать хвост в LOGCAT."
        ],
    )

    add_heading(document, "Итог")
    add_paragraph(
        document,
        "Логи здесь — это отдельный диагностический слой. В файле лежит вся техническая картина "
        "по конкретной сессии: железо, ARCore, конфигурация, runtime-состояния, действия пользователя, "
        "исключения, logcat процесса и итоговый файл для отправки наружу. Если нужно разбирать баг, "
        "этот файл — первая точка входа."
    )

    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
