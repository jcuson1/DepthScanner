# -*- coding: utf-8 -*-
"""Технический документ по режиму Octomap 3D + 3DVFH+ со ссылками на статью и код."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

OUT = Path(r"E:\project\hello_ar_java\docs\Octomap_3DVFH_Технический_документ.docx")


def set_default_font(doc, name="Arial", size_pt=11):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)


def add_title(doc, text, size=20, center=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)


def add_h1(doc, text):
    add_title(doc, text, size=16, center=False)


def add_h2(doc, text):
    add_title(doc, text, size=13, center=False)


def add_h3(doc, text):
    add_title(doc, text, size=12, center=False)


def add_p(doc, text):
    doc.add_paragraph(text)


def add_b(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_n(doc, text):
    doc.add_paragraph(text, style="List Number")


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20)
    r = p.add_run(text)
    r.italic = True


def main():
    doc = Document()
    set_default_font(doc)

    # ──────────────────────────────────────────────────────────────────────
    add_title(doc, "Octomap 3D + 3DVFH+")
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("Технический документ: алгоритмы, код, привязка к научной статье")
    r.italic = True
    r.font.size = Pt(11)

    # ──────────────────────────────────────────────────────────────────────
    add_h1(doc, "1. Источник алгоритма")
    add_p(
        doc,
        "Режим реализует алгоритм 3DVFH+ из научной статьи:",
    )
    add_quote(
        doc,
        "Vanneste S., Bellekens B., Weyn M. «3DVFH+: Real-Time Three-Dimensional "
        "Obstacle Avoidance Using an Octomap». In Modelling and Simulation for "
        "Autonomous Systems (MORSE), 2014, pp. 91–102.",
    )
    add_p(
        doc,
        "Оригинал предназначен для дронов и наземных роботов с RGBD-SLAM-картой и "
        "целью движения. У нас он адаптирован для пешеходного assistive-сценария с "
        "ARCore вместо SLAM-системы. Основное отступление от статьи — этап 3 "
        "(Physical Characteristics: turning circle, climbing motion) опущен, потому "
        "что пешеход разворачивается на месте и не имеет инерции.",
    )

    # ──────────────────────────────────────────────────────────────────────
    add_h1(doc, "2. Структура файлов")
    add_p(doc, "Все файлы режима лежат в пакете:")
    add_code(doc, "app/src/main/java/.../helloar/octomap/")
    add_b(doc, "Octomap.java — разреженная воксельная сетка (HashMap)")
    add_b(doc, "Vfh3DPlanner.java — алгоритм 3DVFH+ (5 этапов)")
    add_b(doc, "OctomapModule.java — оркестратор: depth → октомап → planner")
    add_b(doc, "OctomapInstancedRenderer.java — GPU instanced рендерер кубов")
    add_p(doc, "Внешние зависимости:")
    add_b(
        doc,
        "assets/shaders/octomap_voxel.vert — vertex shader для instanced cubes",
    )
    add_b(
        doc,
        "assets/shaders/floor_heightmap.frag — общий fragment shader (раскраска по высоте)",
    )
    add_b(
        doc,
        "settings/DepthVisualizationSettings.java — enum-значение OCTOMAP_3D + флаг raw/full depth",
    )
    add_b(
        doc,
        "MainActivity.java — wiring: init, onDrawFrame, draw, reset",
    )

    # ──────────────────────────────────────────────────────────────────────
    add_h1(doc, "3. Pipeline режима")
    add_p(
        doc,
        "На каждом кадре ARCore (~30 fps) выполняется следующая последовательность "
        "действий. Все этапы происходят последовательно на GL-thread'е, без "
        "параллельных потоков — простота отладки важнее микро-оптимизаций при текущем "
        "бюджете 5–10 ms на кадр.",
    )

    add_h2(doc, "Шаг A. Получение depth + confidence от ARCore")
    add_b(
        doc,
        "По выбору пользователя в меню — frame.acquireRawDepthImage16Bits() (по умолчанию) "
        "или frame.acquireDepthImage16Bits() (full depth, более плотный, но сглаженный).",
    )
    add_b(doc, "Параллельно — frame.acquireRawDepthConfidenceImage() для фильтра уверенности.")
    add_b(doc, "Код: MainActivity.onDrawFrame, секция if (isOctomapModuleActive) { ... }.")

    add_h2(doc, "Шаг B. Интеграция depth в octomap")
    add_p(doc, "Метод OctomapModule.integrateDepth. Для каждого pixel-а с шагом 2:")
    add_n(
        doc,
        "Median 3×3 (sampleMedianDepthMm) — отсекает одиночные шумовые пики и "
        "edge-артефакты. Если валидных соседей < 5 — пиксель отбрасывается.",
    )
    add_n(
        doc,
        "Confidence filter: если уверенность ARCore < адаптивный порог "
        "(0.20 в ярком свете → 0.45 в темноте) — отбрасываем.",
    )
    add_n(
        doc,
        "Unprojection через intrinsics: optX = (x − cx)·d/fx, optY = (y − cy)·d/fy, "
        "потом ARCore-флипы (Y и Z), потом pose.transformPoint(...) → мировые координаты.",
    )
    add_n(
        doc,
        "rayCastDecrement(...) — 3D-DDA от камеры до точки попадания, decrement каждой "
        "пройденной ячейки через Octomap.integrateMissCell. Если встречен подтверждённый "
        "воксель (hits ≥ 4) — траверс прекращается без decrement'а.",
    )
    add_n(
        doc,
        "Octomap.integrateHit(x, y, z, c) — инкремент endpoint'a. Накапливает: hits, "
        "lastSeenFrame, EMA avgConfidence, Σc, Σc·y.",
    )
    add_p(
        doc,
        "Параметр INTEGRATE_EVERY_NTH_FRAME=2 означает что интеграция выполняется не "
        "каждый кадр, а через один — что даёт ту же среднюю нагрузку CPU при step=2 "
        "и step=4 каждый кадр, но воксели достигают порога подтверждения за половину "
        "времени.",
    )

    add_h2(doc, "Шаг C. Maintenance: decay и cull")
    add_b(
        doc,
        "Раз в DECAY_INTERVAL_FRAMES=30 кадров (~1 сек) — Octomap.decayStaleVoxels: "
        "воксели не наблюдавшиеся ≥ DECAY_STALE_FRAMES=90 кадров (3 сек) теряют 1 hit. "
        "При hits ≤ 0 удаляются. Так чистится «память» о прошлых положениях при "
        "ARCore-дрейфе.",
    )
    add_b(
        doc,
        "Раз в CULL_INTERVAL_FRAMES=60 кадров (~2 сек) — Octomap.cullFarVoxels: "
        "удаление всех вокселей за пределами CULL_RADIUS_M=7 м от пользователя. Ray-casting "
        "до них всё равно не доходит, они только засоряют карту и render mesh.",
    )

    add_h2(doc, "Шаг D. Movement / scanning detection")
    add_b(
        doc,
        "OctomapModule.updateMovementEstimate — EMA вектор движения по разнице позиций "
        "ARCore. Используется как target в planner'е если скорость ≥ 0.25 м/с (иначе "
        "fallback на camera-forward).",
    )
    add_b(
        doc,
        "OctomapModule.updateScanningState — угловая скорость forward-вектора между "
        "кадрами. При ω > 50°/с — режим scanning. Выход из scanning требует затишья "
        "длительностью STABILITY_REQUIRED_NS=500 ms. Во время scanning planner не "
        "вызывается, дирижёр молчит. Один раз озвучивается «Сканирую.» при входе.",
    )

    add_h2(doc, "Шаг E. Запуск 3DVFH+ planner'а (4 Гц)")
    add_b(
        doc,
        "Планировщик вызывается раз в 250 ms (PLANNER_INTERVAL_NS). Чаще не нужно — "
        "голосовая команда длится 1–2 сек, перебивать имеет смысл только при "
        "изменении ситуации.",
    )
    add_b(
        doc,
        "Vfh3DPlanner.plan возвращает Decision { hasPath, azimuthDeg, elevationDeg, "
        "message, severity }, который OctomapModule.buildScanResult оборачивает в "
        "Depth20ScanResult и шлёт дирижёру.",
    )

    add_h2(doc, "Шаг F. Рендер")
    add_b(
        doc,
        "OctomapInstancedRenderer.update собирает per-instance буфер (cx, cy, cz, "
        "weightedAvgY) с frustum culling'ом. Перебор только если octomap.version() "
        "изменился, либо прошло MAX_FRAMES_BETWEEN_REBUILDS=6 кадров.",
    )
    add_b(
        doc,
        "OctomapInstancedRenderer.draw — один glDrawElementsInstanced. Все кубики "
        "рисуются за один draw call.",
    )

    # ──────────────────────────────────────────────────────────────────────
    add_h1(doc, "4. 3DVFH+ алгоритм: статья ↔ код")

    add_h2(doc, "Этап 1 (§4.1 статьи): Octomap Exploring")
    add_p(
        doc,
        "В оригинале — bounding box w_s × w_s × w_s вокруг VCP (Vehicle Center Point). "
        "Проверяются только воксели внутри. У нас bounding sphere радиуса 4 м (это "
        "проще и работает не хуже для пешехода).",
    )
    add_p(doc, "Реализация — внутри цикла Vfh3DPlanner.plan:")
    add_code(
        doc,
        "if (distSq > BOUNDING_RADIUS_M * BOUNDING_RADIUS_M) continue;",
    )

    add_h2(doc, "Этап 2 (§4.2 статьи): 2D Primary Polar Histogram")
    add_p(
        doc,
        "Каждый «занятый» воксель проецируется на 2D полярную гистограмму (азимут × "
        "элевация) с расширением и весом, зависящим от расстояния и уверенности.",
    )

    add_h3(doc, "Уравнение (1) — азимут")
    add_code(doc, "β_z = ⌊(1/α) · arctan((x_i − x_0) / (y_i − y_0))⌋")
    add_p(
        doc,
        "Дискретизация направления на воксель в горизонтальной плоскости. У нас α = 10°, "
        "AZIMUTH_BINS = 36 (полный круг). Реализовано в Vfh3DPlanner.azimuthBin:",
    )
    add_code(
        doc,
        "int bin = (int) Math.floor(azimuthRad / ALPHA_RAD) + FORWARD_AZIMUTH_BIN;\n"
        "return ((bin % AZIMUTH_BINS) + AZIMUTH_BINS) % AZIMUTH_BINS;",
    )

    add_h3(doc, "Уравнение (2) — элевация")
    add_code(doc, "β_e = ⌊(1/α) · arctan((z_i − z_0) / p)⌋,    p = √((x−x₀)² + (y−y₀)²)")
    add_p(
        doc,
        "Угол подъёма воксела относительно горизонта камеры. У нас 7 бинов на ±30° "
        "(Vfh3DPlanner.elevationBin). В оригинальной статье диапазон полная сфера, мы "
        "урезали для пешехода — он не лазит по стенам.",
    )

    add_h3(doc, "Уравнение (4) — footprint enlargement")
    add_code(doc, "λ_ijk = ⌊(1/α) · arcsin((r_r + s + r_v) / d_ijk)⌋")
    add_p(
        doc,
        "Каждый воксель «расширяется» на λ ячеек в каждую сторону, чтобы учесть "
        "размер тела. r_r = USER_RADIUS_M = 0.15 м (полная ширина 30 см), s = "
        "SAFETY_RADIUS_M = 0.20 м (запас), r_v = cellHalf = 0.05 м.",
    )
    add_p(doc, "В коде:")
    add_code(
        doc,
        "float envelope = USER_RADIUS_M + SAFETY_RADIUS_M + cellHalf;\n"
        "float ratio = envelope / dist;\n"
        "float lambda = (float) Math.asin(ratio);\n"
        "int lambdaBins = (int) Math.ceil(lambda / ALPHA_RAD);",
    )

    add_h3(doc, "Уравнение (5) — минимальное расстояние")
    add_code(doc, "l_ijk = d_ijk − (r_r + s + r_v)")
    add_p(
        doc,
        "Расстояние от ПОВЕРХНОСТИ тела пользователя до вокселя (не от центра). Если "
        "воксель внутри тела (l ≤ 0), используется 0.",
    )
    add_code(doc, "float l = Math.max(0f, dist - envelope);")

    add_h3(doc, "Уравнение (6) — вес препятствия в ячейке гистограммы")
    add_code(doc, "H^p_(β_z, β_e) = Σ (o_ijk)² · (a − b · l_ijk),    при e в [β_e ± λ/α], z в [β_z ± λ/α]")
    add_p(
        doc,
        "Каждый воксель вносит вклад во все ячейки в его теневом «прямоугольнике» 2λ×2λ. "
        "Вклад тем больше, чем выше occupancy (= hits / MAX_HITS) и чем ближе воксель.",
    )
    add_p(doc, "В коде:")
    add_code(
        doc,
        "float occ = Math.min(1f, v.hits / (float) Octomap.MAX_HITS);\n"
        "float baseWeight = occ * occ * Math.max(0f, A_CONST - B_CONST * l);",
    )

    add_h3(doc, "Уравнение (7) — нормализация констант a, b")
    add_code(doc, "a − b · ((w_s − 1) / 2)² = 1")
    add_p(
        doc,
        "В статье так выбраны a и b, чтобы при максимальной дальности базовый вес был "
        "ровно 1. У нас линейный вариант: A_CONST = 2, B_CONST = A/BOUNDING_RADIUS. "
        "Это даёт baseWeight = 0 на краю bounding sphere и максимум 2·occ² впритык.",
    )

    add_h2(doc, "Этап 3 (§4.3 статьи): Physical Characteristics — ОПУЩЕН")
    add_p(
        doc,
        "В статье этот этап учитывает turning circle (минимальный радиус разворота "
        "робота) и climbing constant (как ускорение влияет на следующий кадр). Для "
        "пешехода с телефоном эти ограничения не применимы:",
    )
    add_b(doc, "Турнинг-radius у человека ≈ 0 (можно развернуться на месте).")
    add_b(doc, "Подъём по высоте — пешеход всегда идёт по полу.")
    add_p(doc, "Поэтому в код этот этап не вошёл — без потери применимости.")

    add_h2(doc, "Этап 4 (§4.4 статьи): 2D Binary Polar Histogram (eq 18)")
    add_p(
        doc,
        "Преобразование primary гистограммы в булевую через гистерезис с двумя "
        "порогами:",
    )
    add_code(
        doc,
        "if w >= τ_high → binary = true (точно занято)\n"
        "if w <  τ_low  → binary = false (точно свободно)\n"
        "иначе          → сохраняем значение прошлого кадра  (гистерезис)",
    )
    add_p(
        doc,
        "У нас τ_low = 0.5, τ_high = 1.5. Гистерезис подавляет мерцание ячеек "
        "у пограничного веса — без него совет менялся бы каждый кадр.",
    )

    add_h2(doc, "Этап 5 (§4.5 статьи): Path Detection и Selection (eq 19)")
    add_p(doc, "Поиск свободного «коридора» в бинарной гистограмме.")

    add_h3(doc, "Moving window")
    add_p(
        doc,
        "Кандидат-ячейка считается «проходимой» только если все ячейки в квадрате "
        "(2·WINDOW_HALF+1)² вокруг неё тоже свободны. Для тела пользователя нужен "
        "не просто свободный пиксель, а целый коридор.",
    )
    add_code(
        doc,
        "private boolean isPathOpen(int azCenter, int elCenter) {\n"
        "  for (da, de от -WINDOW_HALF до +WINDOW_HALF):\n"
        "    if (binary[azCenter+da][elCenter+de]) return false;\n"
        "  return true;\n"
        "}",
    )

    add_h3(doc, "Уравнение (19) — cost-функция")
    add_code(
        doc,
        "k = μ₁ · Δ(v, k_t) + μ₂ · Δ(v, θ/α) + μ₃ · Δ(v, k_{i-1})",
    )
    add_p(doc, "У нас:")
    add_b(doc, "μ_GOAL = 5.0 — притяжение к цели (target heading) — vehicle движение.")
    add_b(doc, "μ_ELEVATION = 4.0 — штраф за отклонение от уровня глаз.")
    add_b(doc, "μ_SMOOTH = 2.0 — штраф за резкое изменение совета относительно прошлого.")
    add_p(doc, "Реализация — minimum по cost'у среди всех проходимых кандидатов:")
    add_code(
        doc,
        "for каждой пары (a, e):\n"
        "  if !isPathOpen(a, e) skip;\n"
        "  cost = MU_GOAL · azimuthDist(a, FORWARD)\n"
        "       + MU_ELEVATION · |e − LEVEL|\n"
        "       + MU_SMOOTH · (azimuthDist(a, prevA) + |e − prevE|);\n"
        "  if cost < bestCost: bestCost = cost; bestAz = a; bestEl = e;",
    )

    # ──────────────────────────────────────────────────────────────────────
    add_h1(doc, "5. Структуры данных")

    add_h2(doc, "Octomap (Octomap.java)")
    add_p(
        doc,
        "Разреженная воксельная сетка. В отличие от настоящего octree из библиотеки "
        "octomap, мы используем плоский HashMap по упакованному 60-битному ключу. Для "
        "нашего сценария — короткая сессия в одной комнате с цапом 7 000 ячеек — это "
        "работает быстрее.",
    )
    add_p(doc, "Поля Voxel:")
    add_b(doc, "cx, cy, cz — индексы ячейки (int, упакованы в long-ключ).")
    add_b(
        doc,
        "hits — счётчик подтверждений, capped на 30. Растёт с каждым integrateHit, "
        "падает на 1 при integrateMissCell, на 1 каждые 3 сек при decay.",
    )
    add_b(
        doc,
        "lastSeenFrame — номер кадра последнего hit. Используется decay'ем.",
    )
    add_b(
        doc,
        "avgConfidence — EMA уверенности (α=0.3). Используется как фильтр в render и planner.",
    )
    add_b(
        doc,
        "sumWeight (= Σc), sumWeightedY (= Σc·y) — для взвешенного среднего реальной "
        "высоты внутри ячейки. Формула: weightedAvgY = Σc·y / Σc. Это даёт sub-cell "
        "точность по вертикали.",
    )
    add_b(
        doc,
        "version — счётчик изменений в карте. Renderer им batch'ит rebuild'ы.",
    )

    add_h2(doc, "Vfh3DPlanner (Vfh3DPlanner.java)")
    add_b(
        doc,
        "primary[36][7] — float, primary polar histogram. Очищается каждый plan().",
    )
    add_b(
        doc,
        "binary[36][7] — boolean, бинарная гистограмма. СОХРАНЯЕТСЯ между кадрами для "
        "гистерезиса.",
    )
    add_b(
        doc,
        "previousAzimuthBin, previousElevationBin — последний выбор для smoothing'а в eq 19.",
    )

    # ──────────────────────────────────────────────────────────────────────
    add_h1(doc, "6. Параметры (все константы в одном месте)")

    add_h2(doc, "Octomap")
    add_b(doc, "CELL_SIZE_METERS = 0.10 — размер ячейки 10 cm.")
    add_b(doc, "MAX_VOXELS = 7000 — cap.")
    add_b(doc, "MAX_HITS = 30 — потолок счётчика.")
    add_b(doc, "RAYCAST_BLOCKED_THRESHOLD = 4 — выше этого ray-cast не «продырывает».")

    add_h2(doc, "OctomapModule")
    add_b(doc, "DEPTH_PIXEL_STEP = 2, INTEGRATE_EVERY_NTH_FRAME = 2 — частота интеграции.")
    add_b(doc, "MIN_VALID_DEPTH = 0.30 м, MAX_VALID_DEPTH = 5.0 м.")
    add_b(doc, "MIN_CONFIDENCE_BRIGHT/DARK = 0.20 / 0.45 — адаптивный порог.")
    add_b(doc, "MIN_VOXEL_AVG_CONFIDENCE = 0.30 — фильтр воксeля.")
    add_b(doc, "MAX_RAYCAST_STEPS = 50.")
    add_b(doc, "PLANNER_INTERVAL_NS = 250 ms — 4 Гц.")
    add_b(
        doc,
        "DECAY_INTERVAL_FRAMES / STALE_FRAMES / AMOUNT = 30 / 90 / 1 — медленный decay.",
    )
    add_b(doc, "CULL_RADIUS_M = 7.0 м, CULL_INTERVAL_FRAMES = 60 (2 сек).")
    add_b(doc, "MIN_MOVE_SPEED_MPS = 0.25 м/с, MOVE_SMOOTH_ALPHA = 0.3.")
    add_b(
        doc,
        "SCAN_ANGULAR_VELOCITY_RADS = 50°/с, STABILITY_REQUIRED_NS = 500 ms.",
    )

    add_h2(doc, "Vfh3DPlanner")
    add_b(doc, "AZIMUTH_BINS = 36, ELEVATION_BINS = 7 (±30°).")
    add_b(doc, "ALPHA = 10° (= шаг бина).")
    add_b(doc, "BOUNDING_RADIUS_M = 4.0 м.")
    add_b(doc, "USER_RADIUS_M = 0.15 м (ширина 30 см), SAFETY_RADIUS_M = 0.20 м.")
    add_b(doc, "A_CONST = 2.0, B_CONST = A / R.")
    add_b(doc, "TAU_LOW = 0.5, TAU_HIGH = 1.5 — гистерезис.")
    add_b(doc, "MU_GOAL = 5.0, MU_ELEVATION = 4.0, MU_SMOOTH = 2.0.")
    add_b(doc, "OCCUPIED_HIT_THRESHOLD = 6, OCCUPIED_AVG_CONFIDENCE_THRESHOLD = 0.30.")
    add_b(doc, "WINDOW_HALF = 2 — размер коридора в path detection'е.")

    add_h2(doc, "OctomapInstancedRenderer")
    add_b(doc, "RENDER_HIT_THRESHOLD = 6 — на render идут воксели с hits ≥ 6.")
    add_b(doc, "MAX_FRAMES_BETWEEN_REBUILDS = 6 — safety net для batch update'а.")

    # ──────────────────────────────────────────────────────────────────────
    add_h1(doc, "7. Конкурентность и производительность")
    add_p(
        doc,
        "Все этапы выполняются ПОСЛЕДОВАТЕЛЬНО на одном GL thread'е. Параллелизма "
        "между depth integration, planner'ом и render'ом нет — по причинам:",
    )
    add_n(doc, "Простота отладки. Race condition'ы в octomap'е было бы тяжело искать.")
    add_n(
        doc,
        "Нет необходимости — общая нагрузка ~5–10 ms на кадр (бюджет 33 ms при 30 fps), "
        "запас более чем двукратный.",
    )
    add_n(
        doc,
        "ARCore сам блокирует frame.update() если предыдущий кадр не обработан, что "
        "автоматически throttle'ит pipeline.",
    )
    add_p(doc, "Типичные тайминги на Pixel-классе устройств:")
    add_b(doc, "Median + confidence + unproject + ray-cast + hit: 2-4 ms.")
    add_b(doc, "Decay/cull: 0.5-1 ms (раз в секунду).")
    add_b(doc, "Vfh3DPlanner.plan: 1-2 ms (4 раза в сек).")
    add_b(doc, "OctomapInstancedRenderer.update: 0.5-1 ms.")
    add_b(doc, "Draw call: 0.2-0.5 ms на GPU.")

    # ──────────────────────────────────────────────────────────────────────
    add_h1(doc, "8. Голосовые подсказки")
    add_p(
        doc,
        "Vfh3DPlanner.buildPhrase выдаёт фразу на основе |azimuthDeg|:",
    )
    add_b(doc, "< 10°  → «Путь свободен.»             (severity CLEAR)")
    add_b(doc, "10-30° → «Чуть {влево|вправо}.»       (severity INFO)")
    add_b(doc, "30-60° → «Поверните {...}.»            (severity WARNING)")
    add_b(doc, "60-120°→ «Резко {...}.»                (severity STOP)")
    add_b(doc, ">120°  → «Развернитесь {...}.»         (severity STOP)")
    add_b(doc, "При scanning → «Сканирую.» (один раз на эпизод).")
    add_b(doc, "Карта пуста → «Карта пуста. Поводите телефоном вокруг.»")
    add_b(doc, "Все направления заблокированы → «Стоп. Препятствия со всех сторон.»")
    add_p(
        doc,
        "Severity влияет на кулдаун дирижёра (TemporaryConductorModule): STOP=1с, "
        "WARNING=2.5с, INFO=4с, CLEAR=8с.",
    )

    # ──────────────────────────────────────────────────────────────────────
    add_h1(doc, "9. Шейдеры")

    add_h2(doc, "octomap_voxel.vert")
    add_p(
        doc,
        "Vertex shader для GPU instancing'а. Принимает per-vertex unit-cube position "
        "(location 0) и per-instance vec4 (location 1). Все 8 вершин одного куба "
        "получают один a_Instance благодаря glVertexAttribDivisor(1, 1).",
    )
    add_code(
        doc,
        "vec3 worldPos = a_Instance.xyz + a_LocalPos * u_VoxelSize;\n"
        "v_Height = a_Instance.w;\n"
        "v_Valid = 1.0;\n"
        "gl_Position = u_ViewProjection * vec4(worldPos, 1.0);",
    )

    add_h2(doc, "floor_heightmap.frag (общий с floor heightmap mode)")
    add_p(
        doc,
        "Фрагментный шейдер раскрашивает пиксель по {@code v_Height} в палитру: "
        "deep blue → cyan → green → yellow → red. Шкала натягивается на "
        "{@code u_MinHeight..u_MaxHeight} (динамические).",
    )

    # ──────────────────────────────────────────────────────────────────────
    add_h1(doc, "10. Известные ограничения")
    add_b(
        doc,
        "ARCore-дрейф: после ~30 сек ходьбы карта может «двоиться» на удалённых "
        "областях. Decay компенсирует, но не идеально. Полное решение — loop closure "
        "SLAM, в ARCore не поддерживается.",
    )
    add_b(
        doc,
        "Зависимость от качества depth: блики, тёмные углы и зеркала дают шум. Adaptive "
        "confidence threshold подавляет, но не устраняет полностью.",
    )
    add_b(
        doc,
        "Алгоритм не различает тип препятствия (стол, человек, стена). Только "
        "геометрия. Для семантики нужен YOLO/CNN-детектор поверх.",
    )
    add_b(
        doc,
        "Динамические объекты (идущий человек) попадают в карту, ray-casting удалит "
        "за 1-2 сек. Если объект задержался на одном месте на 5+ сек — он становится "
        "«статичным» в восприятии алгоритма.",
    )
    add_b(
        doc,
        "Нет глобального планирования (доведи до точки B). Алгоритм даёт только "
        "локальный совет «куда идти прямо сейчас». Для full-route navigation нужна "
        "карта помещения + Dijkstra/A*.",
    )

    # ──────────────────────────────────────────────────────────────────────
    add_h1(doc, "11. Ссылки на код (быстрая навигация)")
    add_b(doc, "Octomap.integrateHit() — Octomap.java")
    add_b(doc, "Octomap.integrateMissCell() — Octomap.java")
    add_b(doc, "Octomap.decayStaleVoxels(), cullFarVoxels() — Octomap.java")
    add_b(doc, "OctomapModule.process() — главный pipeline frame'а")
    add_b(doc, "OctomapModule.integrateDepth() — depth → octomap")
    add_b(doc, "OctomapModule.rayCastDecrement() — 3D-DDA")
    add_b(doc, "OctomapModule.sampleMedianDepthMm() — median 3×3")
    add_b(doc, "OctomapModule.computeEffectiveMinConfidence() — adaptive по свету")
    add_b(doc, "OctomapModule.updateMovementEstimate() — EMA вектор движения")
    add_b(doc, "OctomapModule.updateScanningState() — детектор сканирования")
    add_b(doc, "Vfh3DPlanner.plan() — все 5 этапов 3DVFH+")
    add_b(doc, "Vfh3DPlanner.isPathOpen() — moving window check")
    add_b(doc, "Vfh3DPlanner.buildPhrase() — голосовая фраза")
    add_b(doc, "OctomapInstancedRenderer — GPU instanced cubes")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved to {OUT}")


if __name__ == "__main__":
    main()
